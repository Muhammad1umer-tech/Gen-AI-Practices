# InternetSearchAnalyst (go thru different urls and create summary and give to InsightRearcher)
# InsightResearcher (indept search on each topic and will give use detailed report)
# duckduckgo search and beautiful soup
#Tools (search the internet, and read) duckduckgo_search, beautiful soup

# supervisor
from pydantic import BaseModel, Field
from langchain_core.tools import tool
from typing import Optional, Literal
from duckduckgo_search import DDGS
from typing import List, Dict
from bs4 import BeautifulSoup
from docx import Document
import urllib.parse
import requests
import json


@tool
def InternetSearchTool(query: str) -> list:
    '''This tool take the query to do web search on it and
    return the results along with title and link'''
    
    print("\nInternet search tool called", query)
    
    total_links = 1
    results = []
    
    # Step 1: Define the DuckDuckGo search URL and search query
    encoded_query = urllib.parse.quote(query)
    url = f"https://html.duckduckgo.com/html/?q={encoded_query}"

    # Step 2: Send a request to DuckDuckGo
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/92.0.4515.159 Safari/537.36"
    }
    response = requests.get(url, headers=headers)

    # Step 3: Check if the request was successful
    if response.status_code == 200:
        # Step 4: Parse the HTML with Beautiful Soup
        soup = BeautifulSoup(response.text, "html.parser")

        # Step 5: Extract search results
        for link in soup.find_all('a', class_="result__a"):
            title = link.get_text()
            href = link['href']
            # Collect result details
            results.append({
                "title": title,
                "link": href
            })

            if len(results) == total_links:
                break

    else:
        print("Failed to retrieve results from DuckDuckGo.")
        
    print(results)
    return results # [{title, link}]

class DataModel(BaseModel):
    title: str
    link: str

class InsightResearcherArgs(BaseModel):
    datas: List[DataModel]



@tool("InsightResearcher", args_schema=InsightResearcherArgs, return_direct=True)
def InsightResearcher(datas: List[Dict[str, str]]):
    
    """This tool takes a list of DataModel instances, each containing a title and a link.
    It fetches the content at each URL and returns the extracted content.
    """
    
    # datas = json.loads(arg)
    print("\nmessages InsightResearcher ", datas)
    results = []    
    for data in datas:
        response = requests.get(data.link)
        soup = BeautifulSoup(response.content, 'html.parser')
        results.append({
            'text': data.title,
            'content': soup.get_text()})
            # 'content': "HI, web scrap"})
    
    print("results from InsightResearcher ", results)
    return results


class DataModelDocumentWriter(BaseModel):
    title: str
    content: str

class DocWriterArgs(BaseModel):
    datas: List[DataModelDocumentWriter]

@tool("DocumentWriterTool", args_schema=DocWriterArgs, return_direct=True)
def DocumentWriterTool(datas:  List[Dict[str, str]]) -> str:
    """
    This tools creates a Word document for each entry in a provided list of DataModelDocumentWriter.

    Args:
        This tool takes a list of DataModelDocumentWriter instances, each containing a title and a content.

    Returns:
        str: A success message indicating the documents were saved, or an error message.
    """
    print("\nMessages received by DocumentWriter: ", datas)
    
    try:
        for data in datas:
            title = data.title
            content = data.content

            # Create a new Document
            doc = Document()
            doc.add_heading(title, 0)
            doc.add_paragraph(content)
            
        doc.save("Document.docx")
        
        print("\nEnd of DocumentWriter:")
        return "Successfully saved the documents"

    except Exception as e:
        error_message = f"An error occurred while saving the documents: {e}"
        print(error_message)
        return error_message

# print(InsightResearcher({
#     "datas": InternetSearchTool("Web scrapping")
# }))